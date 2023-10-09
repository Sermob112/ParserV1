import csv

from bs4 import BeautifulSoup
import requests
import locale
import re

import json
import os
from selenium import webdriver
from docx import Document

class  ParserOld:
    def __init__(self):
        pass

    def makeLinkNum(self, numer,filePath):
        self.filePath = filePath
    # num = numer
        try:
            search_url = f'https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString={numer}&morphology=on&search-filter=Дате+размещения&pageNumber=1&sortDirection=false&recordsPerPage=_10&showLotsInfoHidden=false&sortBy=UPDATE_DATE&fz44=on&fz223=on&af=on&ca=on&pc=on&pa=on&currencyIdGeneral=-1'
            # test_url3 = f'https://zakupki.gov.ru/epz/order/notice/ok20/view/common-info.html?regNumber={numer}'
            HEADERS_test = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.0.0'
                                ' YaBrowser/23.3.0.2246 Yowser/2.5 Safari/537.36', 'accept': '*/*'}
            req = requests.get(search_url, headers=HEADERS_test, params=None)
            src = req.text
            soup = BeautifulSoup(src, 'lxml')
            status = 'Успешное подключение'
            col = soup.find('div', class_ = 'registry-entry__header-mid__number')
            a_tag = col.find('a')
            match = re.search(r'noticeInfoId=(\d+)', a_tag['href'])
            number = match.group(1)
            self.main_directory = 'Закупка № ' + str(numer)
            self.num = numer
            self.agent(numer = number)
            # return  number
        except:
            status = 'Ошибка подключения'
    


    def agent(self, numer):
        # self.num = numer
        try:
            # test_url3 = f'https://zakupki.gov.ru/epz/order/notice/ok20/view/common-info.html?regNumber={numer}'
            test_url2 = f'https://zakupki.gov.ru/epz/order/notice/notice223/common-info.html?noticeInfoId={numer}'
            self.HEADERS = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.0.0'
                              ' YaBrowser/23.3.0.2246 Yowser/2.5 Safari/537.36', 'accept': '*/*'}

            req = requests.get(test_url2, headers=self.HEADERS, params=None)
            src = req.text
            self.soup = BeautifulSoup(src, 'lxml')
            self.status = 'Успешное подключение'
            return self.soup
        except:
            self.status = 'Ошибка подключения'

    ##############################################################################################################
    #основная информация
    ######################################################
    #Серийный номер

    def parse_head(self ):
        mass = []
        serial_date = {}

        source = self.soup.find(class_="col-6 pr-0 mr-21px")
        # self.ObjectName =[]
        self.ObjectName = source.find(class_='registry-entry__body-value').get_text().strip().replace('"','')
        # reg_num = self.soup.find(class_='registry-entry__header-mid__number').get_text().strip()
        # self.main_directory = 'Закупка ' + str(reg_num)
        lines = source.get_text().strip().splitlines()
        # Удалите пустые строки
        mainMass = [line.strip() for line in lines if line.strip()]
        sourceleft = self.soup.find(class_="col d-flex flex-column registry-entry__right-block b-left")
        lines = sourceleft.get_text().strip().splitlines()
        leftMass = [line.strip() for line in lines if line.strip()]
        cleaned_left = [item.replace('\xa0', ' ') for item in leftMass]

        # return  mainMass+cleaned_left
        return self.ObjectName
    
    def get_links(self):
        tabs_of_links = {}
        link_razdels = self.soup.find(class_='container card-layout')
        if link_razdels != None:
            try:
                link_razdels = self.soup.find(class_ = 'tabsNav d-flex').find_all('a')
                for links in link_razdels:
                    linkl = 'https://zakupki.gov.ru/' + links.get('href')
                    title_razde = links.text.strip()
                    tabs_of_links[title_razde] = linkl
                    
            except:
                link_razdels = self.soup.find(class_='tabsNav d-flex align-items-end').find_all('a')
                for links in link_razdels:
                    linkl = 'https://zakupki.gov.ru/' + links.get('href')
                    title_razde = links.text.strip()
                    tabs_of_links[title_razde] = linkl
            return tabs_of_links

    def mainInfo(self):
        mainMass = []
        newmainMass = []
        JornalMass = []
        tabs_of_links = self.get_links()
        for title, link in tabs_of_links.items():
            req = requests.get(url=link, headers=self.HEADERS)
            src = req.text
            soup = BeautifulSoup(src, "lxml")
            if title == 'Журнал событий':
                JornalMass = self.get_jornals(link)
            if title == 'Документы':
                JornalMass = self.get_documents(link)
            try:
                
                containerMain = soup.find_all(class_='card-common-content')
                containerMain2 = soup.find_all(class_='card-attachments')
            
                if containerMain != None:
                    for i in containerMain:
                        lines = i.get_text().strip().splitlines()
                        Mass = [line.strip() for line in lines if line.strip()]
                        cleaned_Mass = [item.replace('\xa0', ' ') for item in Mass]
                        mainMass =  mainMass + cleaned_Mass
                
                if containerMain2 != None:
                    for i in containerMain2:
                        lines = i.get_text().strip().splitlines()
                        Mass = [line.strip() for line in lines if line.strip()]
                        cleaned_Mass = [item.replace('\xa0', ' ') for item in Mass]
                        mainMass =  mainMass + cleaned_Mass
                        newmainMass =  newmainMass + cleaned_Mass
                
            except:
                print('Error')
        mainMass.append(newmainMass)
        mainMass.append(JornalMass)
        return mainMass
    
    def get_jornals(self, link):
        JornalMass= []
        driver = webdriver.Chrome()
        driver.get(link)
        import time
        time.sleep(0)
        html = driver.page_source
        soup = BeautifulSoup(html, 'lxml')
        driver.quit()
        containerMain3 = soup.find_all(class_='tabBoxWrapper tabBoxWrapper__mb24')
        if containerMain3 != None:
            for i in containerMain3:
                lines = i.get_text().strip().splitlines()
                Mass = [line.strip() for line in lines if line.strip()]
                JornalMass = JornalMass + Mass
        return JornalMass
    
    def get_documents(self, link):
        titlesMass = []
        linkMass = []
        req = requests.get(link, headers=self.HEADERS, params=None)
        src = req.text
        soup = BeautifulSoup(src, 'lxml')
        findTitle = soup.find_all(class_='container card-attachments-container')
        for titles in findTitle:
            title = titles.find(class_='title')
            if title:
                title_text = title.get_text().strip()
                values = titles.find_all(class_='col-6 b-left')
                for laxir in values:
                    luxit = laxir.find_all('a')
                    try:
                        for links in luxit:
                            href = links.get('href')
                            if href and 'download/download.html?id' in href:
                                linkl = href
                                tooltip_value = links.get('data-tooltip')
                                tooltip_soup = BeautifulSoup(tooltip_value, 'html.parser')
                                span_element = tooltip_soup.find('span')
                                if span_element:
                                    text = span_element.get_text(strip=True).replace('"','')
                                    
                                    # linkMass.append(linkl)
                                    response = requests.get(f'https://zakupki.gov.ru{linkl}', headers=self.HEADERS)
                                    if response.status_code == 200:
                                        # path = os.path.join(self.main_directory, text)
                                        # os.makedirs(path)
                                        # Создаем каталог для сохранения файлов, если его нет
                                        os.makedirs(f'{self.filePath}/{self.main_directory + self.ObjectName}/{title_text}', exist_ok=True)
                                        # Сохраняем файл в указанный каталог
                                        with open(f'{self.filePath}/{self.main_directory + self.ObjectName}/{title_text}/{text}', "wb") as f:
                                            f.write(response.content)
                                    else:
                                        print(f"Не удалось скачать файл: {linkl}")
                    except Exception as e:
                        print(f"Произошла ошибка: {str(e)}")

    def makeDoc(self):
        headMass = self.parse_head()
        mainMass= self.mainInfo()
        doc = Document()
        data = headMass + mainMass
        structured_data = {}

        
        # for i in range(0, len(data) - 1, 2):
        #     key = data[i]
        #     value = data[i + 1]
        #     structured_data[key] = value

     
        # if len(data) % 2 != 0:
        #     unstructured_data = data[-1]

        # doc.add_heading('Данные о закупке', level=1)

        
        # for key, value in structured_data.items():
        #     doc.add_paragraph(f'{key}: {value}')

        # doc.add_heading('Неструктурированные данные', level=2)
        # doc.add_paragraph(unstructured_data)

        # Добавляем каждый элемент массива данных в документ
        for item in data:
            doc.add_paragraph(item)
        doc.save(f"{self.filePath}/{self.main_directory}/Все данные о закупке №{self.num}.docx")
        # return data

        





# par = ParserOld()
# par.agent('9272643')
# print(par.parse_head())
# par.makeDoc()

# par.parse_head()
# print(par.get_links())
# print(par.mainInfo())
# print(par.get_jornals('https://zakupki.gov.ru/epz/order/notice/notice223/event-journal.html?noticeInfoId=5963715'))
# print(par.get_documents('https://zakupki.gov.ru/epz/order/notice/notice223/documents.html?noticeInfoId=12821871'))