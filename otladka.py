from typing import List, Dict, Any

from bs4 import BeautifulSoup
import requests
import urllib3
import locale
import re
import csv
from docx import Document
import pandas as pd
import json
import os
# from selenium import webdriver
# link = 'https://zakupki.gov.ru/epz/order/notice/ok20/view/supplier-results.html?regNumber=0172500000423000006'

# from selenium import webdriver
# from selenium.webdriver.common.by import By
# from selenium.webdriver.support.ui import WebDriverWait
# from selenium.webdriver.support import expected_conditions as EC




# driver = webdriver.Chrome()

# # Переходим на страницу
# driver.get("https://example.com")

# # Используем метод find_element_by_class_name для поиска элемента с классом "hidden"
# elem = driver.find_element(By.CLASS_NAME,"tableBlock__col tableBlock__row")

# # Выводим содержимое элемента
# print(elem.text)

# # Закрываем браузер
# driver.quit()

# def supplier_result():
#     block_title = []
#     title_data= []
#     value_data = []
#     list_info= []
#     zakazchic_data= []
#     table_titles = []
#     all_table_info = []
#     doc_titles = []
#     new_dict = {}
#     i = 0
#     n = 0
#     # row_info = soup.find_all(class_= 'row blockInfo')
#     # for info in row_info:
#     driver = webdriver.Firefox()
#     driver.get(link)
#     import time
#     time.sleep(5)
#     html = driver.page_source

#     soup = BeautifulSoup(html, 'lxml')



#     col_9 = soup.find_all(class_='tableBlock__body')
#     print(col_9)


#     driver.quit()

# # supplier_result()

def journal_of_events():

        th_mass = []
        data = []
        fn_date= {}
        i = 0
        test_url3 = f'https://zakupki.gov.ru/epz/order/notice/ok44/view/event-journal.html?regNumber=0373100112518000004'
            
        HEADERS_test = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.0.0'
                              ' YaBrowser/23.3.0.2246 Yowser/2.5 Safari/537.36', 'accept': '*/*'}
        params ={

                'regNumber': '0373100112518000004' 
            }
        req = requests.get(test_url3, headers=HEADERS_test, params=None)
        src = req.text
        soup = BeautifulSoup(src, 'lxml')
        # driver.quit()
        wrapper = soup.find('div', class_='container')
        return soup.get_text()
        #не нашел журнал событий
        try:
            for tbn in wrapper:
                c_th = tbn.find_all('th')
                for th_items in c_th:
                    th_text = th_items.get_text().strip()
                    th_mass.append(th_text)

                c_td = tbn.find_all('td')
                for td_items in c_td:
                    td_text = td_items.get_text().strip()
                    try:
                        data.append({th_mass[i]: td_text})
                    except Exception:
                        i = 0
                        data.append({th_mass[i]: td_text})
                    i += 1
            status = 'Успешный парсинг Журнала'
            return data

        except Exception:
            status = 'Ошибка парснига  Журнала'

# print(journal_of_events())

def makeLinkNum(numer):
    num = numer
    try:
        search_url = f'https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString={numer}&morphology=on&search-filter=Дате+размещения&pageNumber=1&sortDirection=false&recordsPerPage=_10&showLotsInfoHidden=false&sortBy=UPDATE_DATE&fz44=on&fz223=on&af=on&ca=on&pc=on&pa=on&currencyIdGeneral=-1'
        # test_url3 = f'https://zakupki.gov.ru/epz/order/notice/ok20/view/common-info.html?regNumber={numer}'
        HEADERS_test = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.0.0'
                            ' YaBrowser/23.3.0.2246 Yowser/2.5 Safari/537.36', 'accept': '*/*'}
        req = requests.get(search_url, headers=HEADERS_test, params=None)
        src = req.text
        soup = BeautifulSoup(src, 'lxml')
        status = 'Успешное подключение'
        col = soup.find('div', class_ = 'registry-entry__header-mid__number')
        a_tag = col.find('a')
        match = re.search(r'noticeInfoId=(\d+)', a_tag['href'])
        number = match.group(1)
        return  number
    except:
        status = 'Ошибка подключения'


# print(agent('31704961137'))


################################
# data = ['223-ФЗ Прочие', '№ 31704961137', 'Закупка завершена', 'Объект закупки', 'проведение работ и услуг по строительству и поставки многофункционального судна снабжения 001', 'Заказчик', 'АКЦИОНЕРНОЕ ОБЩЕСТВО "ДАЛЬНЕВОСТОЧНЫЙ ЦЕНТР СУДОСТРОЕНИЯ И СУДОРЕМОНТА"', 'Начальная цена', '11 821 240 000,00 ₽', 'Размещено', '31.03.2017', 'Обновлено', '31.03.2017', 'Сведения о закупке', 'Реестровый номер извещения', '31704961137', 'Способ осуществления закупки', 'Внутригрупповая закупка', 'Наименование закупки', 'проведение работ и услуг по строительству и поставки многофункционального судна снабжения 001', 'Редакция', '1', 'Дата размещения извещения', '31.03.2017', 'Дата размещения текущей редакции извещения', '31.03.2017', 'Сведения о заказчике', 'Наименование организации', 'АКЦИОНЕРНОЕ ОБЩЕСТВО "ДАЛЬНЕВОСТОЧНЫЙ ЦЕНТР СУДОСТРОЕНИЯ И СУДОРЕМОНТА"', 'ИНН', '2536196045', 'КПП', '253601001', 'ОГРН', '1072536016211', 'Место нахождения', '690001, Г.. ВЛАДИВОСТОК, УЛ. СВЕТЛАНСКАЯ, Д.72', 'Почтовый адрес', '690001, Приморский край, г. Владивосток, ул. Светланская, д. 72', 'Контактная информация', 'Адрес электронной почты', 'dcss@dcss.ru', 'Контактный телефон', '+8 (423) 2651736', 'Порядок проведения процедуры', 'Рассмотрение заявок', 'Место рассмотрения заявок', 'Владивосток, Светланская, 72', 'Дата рассмотрения заявок', '31.03.2017 (МСК+7)', 'Подведение итогов', 'Место подведения итогов', 'Владивосток, Светланская, 72', 'Дата подведения итогов', '31.03.2017', 'Предоставление документации', 'Срок предоставления', 'с 31.03.2017 по', '31.03.2017', 'Место предоставления', 'не применимо', 'Порядок предоставления', 'не применимо', 'Официальный сайт ЕИС, на котором размещена документация', 'www.zakupki.gov.ru', 'Внесение платы за предоставление конкурсной документации', 'Требования не установлены', 'Сведения о лотах', 'Номер, наименование лота', 'Централизованная закупка', 'Сведения о цене договора', 'Классификация по ОКПД2', 'Классификация по ОКВЭД2', '1', 'Строительство и поставка многофункционального судна снабжения  усиленного ледового класса', 'Нет', 'Начальная (максимальная) цена договора:', '11 821 240 000,00 ₽', '30.11 Корабли, суда и плавучие конструкции', '30.11 Строительство кораблей, судов и плавучих конструкций', '№', 'Критерий', 'Описание', 'Максимальное значение', 'Вес', 'Место поставки товара, выполнения работы, оказания услуги', 'Место поставки товара, выполнения работы, оказания услуги (субъект РФ)', 'Приморский край', 'Место поставки товара, выполнения работы, оказания услуги (адрес)', 'г. Большой Камень', 'Документация по закупке', 'Извещение о закупке № 31704961137 (Версия №1)', 'Размещено', '31.03.2017 09:22', '(МСК+7)', 'Редакция', 'Действующая', 'Прикрепленные файлы', 'Закупочная документация', 'Протоколы работы комиссии', 'Иной протокол №31704961137-01', 'Размещено', '31.03.2017 09:47', '(МСК+7)', 'Редакция', 'Действующая', 'Прикрепленные файлы', 'Протокол', 'Извещение о проведении закупки', 'Наименование документа', 'Редакция', 'Статус', 'Размещено', 'Извещение о закупке № 31704961137', '1', 'Действующая', '31.03.2017 09:22 (МСК+7)', 'Изменения извещения', 'Сведения отсутствуют', 'Изменения извещения с отменой лотов', 'Сведения отсутствуют', 'Изменения извещения со сменой заказчика', 'Сведения отсутствуют', 'Разъяснения положений документации', 'Сведения отсутствуют', 'Отмена закупки', 'Сведения отсутствуют', 'Иной протокол', 'Наименование документа', 'Редакция', 'Статус', 'Размещено', 'Иной протокол №31704961137-01', '1', 'Действующая', '31.03.2017 09:47 (МСК+7)', 'Реестр договоров', 'Наименование документа', 'Редакция', 'Статус', 'Размещено', 'Договор №62536196045170000060003', '2', 'Действующая', '29.01.2019 09:20', '(МСК+7)', 'Договор №62536196045170000060001', '1', 'Недействующая', '02.06.2017 05:30', '(МСК+7)', ['Документация по закупке', 'Извещение о закупке № 31704961137 (Версия №1)', 'Размещено', '31.03.2017 09:22', '(МСК+7)', 'Редакция', 'Действующая', 'Прикрепленные файлы', 'Закупочная документация', 'Протоколы работы комиссии', 'Иной протокол №31704961137-01', 'Размещено', '31.03.2017 09:47', '(МСК+7)', 'Редакция', 'Действующая', 'Прикрепленные файлы', 'Протокол', 'Извещение о проведении закупки', 'Наименование документа', 'Редакция', 'Статус', 'Размещено', 'Извещение о закупке № 31704961137', '1', 'Действующая', '31.03.2017 09:22 (МСК+7)', 'Изменения извещения', 'Сведения отсутствуют', 'Изменения извещения с отменой лотов', 'Сведения отсутствуют', 'Изменения извещения со сменой заказчика', 'Сведения отсутствуют', 'Разъяснения положений документации', 'Сведения отсутствуют', 'Отмена закупки', 'Сведения отсутствуют'], ['Дата и время', 'Событие', '03.04.2017 05:05 (МСК+7)', 'Размещены первоначальные сведения о договоре «строительство и поставка многофункционального судна 001».', 'Извещение о закупке: № 31704961137 «проведение работ и услуг по строительству и поставки многофункционального судна снабжения 001 ».', 'Заказчик: АКЦИОНЕРНОЕ ОБЩЕСТВО "ДАЛЬНЕВОСТОЧНЫЙ ЦЕНТР СУДОСТРОЕНИЯ И СУДОРЕМОНТА".', 'Организация, разместившая сведения: АКЦИОНЕРНОЕ ОБЩЕСТВО "ДАЛЬНЕВОСТОЧНЫЙ ЦЕНТР СУДОСТРОЕНИЯ И СУДОРЕМОНТА".', '31.03.2017 10:16 (МСК+8)', 'Закупка переведена на этап «Размещение завершено» с этапа «Работа комиссии»', '31.03.2017 10:02 (МСК+8)', 'Закупка переведена на этап «Работа комиссии» с этапа «Размещение завершено»', '31.03.2017 09:47 (МСК+8)', 'Закупка переведена на этап «Размещение завершено» с этапа «Работа комиссии»', '31.03.2017 09:47 (МСК+8)', 'Размещен протокол № 31704961137-01 «Протокол утверждения закупочной документации и подписание договора».', 'Извещение о закупке: № 31704961137 «проведение работ и услуг по строительству и поставки многофункционального судна снабжения 001 ».', 'Способ закупки: Внутригрупповая закупка.', 'Заказчик: АКЦИОНЕРНОЕ ОБЩЕСТВО "ДАЛЬНЕВОСТОЧНЫЙ ЦЕНТР СУДОСТРОЕНИЯ И СУДОРЕМОНТА".', 'Организация, разместившая сведения: АКЦИОНЕРНОЕ ОБЩЕСТВО "ДАЛЬНЕВОСТОЧНЫЙ ЦЕНТР СУДОСТРОЕНИЯ И СУДОРЕМОНТА".', '31.03.2017 09:22 (МСК+8)', 'Закупка переведена на этап «Работа комиссии» с этапа «Подача заявок»', '31.03.2017 09:22 (МСК+8)', 'Закупка переведена на этап «Подача заявок» с этапа «Формирование извещения»', '31.03.2017 09:22 (МСК+8)', 'Размещено извещение о закупке № 31704961137 «проведение работ и услуг по строительству и поставки многофункционального судна снабжения 001 ».', 'Способ закупки: Внутригрупповая закупка.', 'Заказчик: АКЦИОНЕРНОЕ ОБЩЕСТВО "ДАЛЬНЕВОСТОЧНЫЙ ЦЕНТР СУДОСТРОЕНИЯ И СУДОРЕМОНТА".', 'Организация, разместившая сведения: АКЦИОНЕРНОЕ ОБЩЕСТВО "ДАЛЬНЕВОСТОЧНЫЙ ЦЕНТР СУДОСТРОЕНИЯ И СУДОРЕМОНТА".']]
# structured_data = {}

# # Проходим по данным с шагом 2 (чтобы использовать каждый второй элемент как ключ, а следующий - как значение)
# for i in range(0, len(data) - 1, 2):
#     key = data[i]
#     value = data[i + 1]
#     structured_data[key] = value

# # Последний элемент (если он есть) остается неструктурированным
# if len(data) % 2 != 0:
#     unstructured_data = data[-1]

# # Теперь structured_data - это словарь, в котором ключи и значения соответствуют вашим данным
# print(structured_data)

# # Переменная unstructured_data содержит последний элемент (неструктурированный)
# if len(data) % 2 != 0:
#     print("Неструктурированные данные:", unstructured_data)
################################
def makeMass(csv_filename):
    data_from_second_column = []
    # Открываем CSV файл для чтения
    with open(csv_filename, newline='', encoding='cp1251') as csvfile:
        reader = csv.reader(csvfile, delimiter=';')

        # Пропускаем заголовок, если есть
        next(reader, None)

        # Читаем вторую строку и выбираем вторую часть (индекс 1)
        for row in reader:
                # Проверяем, что строка содержит как минимум два элемента (для второго столбца)
                if len(row) >= 2:
                    # Добавляем элемент второго столбца в массив
                    data_from_second_column.append(row[1].replace("№", ""))
    print(data_from_second_column)



def main_info_body(numer):
        test_url3 = f'https://zakupki.gov.ru/epz/order/notice/ok20/view/common-info.html?regNumber={numer}'
        HEADERS = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.0.0'
                              ' YaBrowser/23.3.0.2246 Yowser/2.5 Safari/537.36', 'accept': '*/*'}
        req = requests.get(test_url3, headers=HEADERS, params=None)
        src = req.text
        soup = BeautifulSoup(src, 'lxml')
        status = 'Успешное подключение'
        main_infp = soup.find_all(class_ ='common-text b-bottom pb-3')
        list_info = []
        all_table_info = []
        data_td = []
        block_title = []
        new_dict = {}
        table_info = []
        try:
            if len(main_infp) != 0:
            # title_cap = soup.find(class_= 'common-text__caption').text.strip()
            # print(title_cap)
                for info in main_infp:
                    #весь текст
                    # zagolovok = info.text.strip()
                    col_9 = info.find_all(class_='col-9 mr-auto')
                    title_cap = info.find(class_='common-text__caption').text.strip()
                    list_info.append({title_cap})
                    for values in col_9:
                        # title_cap = info.find(class_='common-text__caption')

                        title_text = values.find(class_= 'common-text__title')
                        value_text = values.find(class_= 'common-text__value')
                        if (title_text == None  or value_text == None):
                            continue
                        else:
                            title_text = values.find(class_='common-text__title').text.strip()
                            value_text = values.find(class_='common-text__value').text.strip()
                        list_info.append({title_text:value_text})


            else:
                b = 0
                n = 0
                main_infp = soup.find_all(class_='row blockInfo')
                for info in main_infp:
                    # col_9 = info.find_all(class_='blockInfo__section section')
                    col_9 = info.find_all(class_='blockInfo__section')
                    title_cap = info.find(class_='blockInfo__title').get_text().strip()
                    block_title.append(title_cap)
                    for values in col_9:
                        title_text = values.find(class_='section__title')
                        value_text = values.find(class_='section__info')
                        if (title_text == None  or value_text == None):
                            continue
                        else:
                            title_text = values.find(class_='section__title').get_text().strip()
                            value_text = values.find(class_='section__info').get_text().strip()
                            if '\n' in title_text or '\n' in value_text:
                                title_text = title_text.replace('\n', '')  # remove non-breaking spaces
                                value_text = value_text.replace('\n', '')
                            if ' ' in value_text:
                                value_text = value_text.replace(' ', '')  # remove non-breaking spaces
                            list_info.append({title_text:value_text})

                    table = info.find_all(class_ ='blockInfo__table tableBlock')
                    if table != None:
                        for tables in table:
                            # table = info.find(class_='blockInfo__table tableBlock')
                            hiegth_row = tables.find_all('th')
                            table_td = tables.find_all('td')

                            for i in hiegth_row:
                                column_name = i.text.strip()
                                all_table_info.append(column_name)

                            for i in table_td:
                                td = i.get_text().strip()
                                if(td == ''):
                                    continue
                                if '\n' in td:
                                    td = td.replace('\n', '') # remove non-breaking spaces
                                    td = re.findall('[A-ZА-Я][^A-ZА-Я]*', td)
                                if '\xa0' in td:
                                    td = td.replace('\xa0', '').replace(',', '.')   # remove non-breaking spaces

                                data_td.append(td)
                                try:
                                    list_info.append({all_table_info[n]:td})
                                except IndexError:
                                    # continue
                                    n = 0
                                    list_info.append({all_table_info[n]: td})
                                n = n + 1
                    new_dict[title_cap] = list_info
                    list_info = []
                    b = b + 1
                # return data_td
            # self.status = 'успешный парсинг Общей информации'
            return new_dict
        except Exception:
            status = 'Общая информация ошибка парсинга'

# print(main_info_body('1020600000221000005'))
            
def collapse_element(numer):
    test_url3 = f'https://zakupki.gov.ru/epz/order/notice/ok20/view/common-info.html?regNumber={numer}'
    HEADERS = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.0.0'
                            ' YaBrowser/23.3.0.2246 Yowser/2.5 Safari/537.36', 'accept': '*/*'}
    req = requests.get(test_url3, headers=HEADERS, params=None)
    src = req.text
    soup = BeautifulSoup(src, 'lxml')
    i = 0
    j = 0
    collaps_titles  = []
    collaps_main_title_mass = []
    cont_main_title_mass = []
    th_mass = []
    titls_table_main = []
    collaps_main = []
    collaps_data = {}
    conteiner_text = []
    table_data = {}
    conteiner_titls_data = {}
    fgh = 0
    collaps_content = soup.find_all( 'div', class_ = 'blockInfo__collapse collapseInfo')
    try:
        if collaps_content != None:
            for infos in collaps_content:
                collaps_title = infos.find( class_ = 'collapse__title_text').get_text()
                collaps_title = collaps_title.replace('\n', '').replace('\xa0', ' ')
                collaps_titles.append(collaps_title)
                content_block_info  = infos.find_all('div', class_="content__block blockInfo")
                collaps_data[0] = collaps_title
                for item in content_block_info:

                    block_info_title = item.find( class_ = 'blockInfo__title').get_text()
                    collaps_titles.append(block_info_title)
                    block_sec_title = item.find_all(class_='section__title')
                    for its in block_sec_title:
                        collaps_main_title_mass.append(its.get_text().replace('\n', ''))
                    block_sec_info = item.find_all(class_='section__info')
                    for it in block_sec_info:
                        # block_sec_info_mass.append(it.get_text().replace('\n', '').replace('\xa0', ' '))
                        try:
                            collaps_main.append(
                                {collaps_main_title_mass[fgh]: it.get_text().replace('\n', '').replace('\xa0', ' ')})
                        except IndexError:
                            collaps_main.append({'': it.get_text().replace('\n', '').replace('\xa0', ' ')})
                        fgh = fgh +1

                    collaps_data[block_info_title] = collaps_main
                    collaps_main = []
                #####парсинг таблицы
                conteiner_table= infos.find_all('div', class_="container")
                if conteiner_table != None:
                    conteiner_table = infos.find_all('div', class_="container")

                    ###################################################
                    #Парсинг значений таблицы
                    for tb in conteiner_table:
                        # может быть несколько тайтлов\ пока только один
                        row_block_info = tb.find_all(class_ = 'row blockInfo')
                        for rower in row_block_info:
                            tit = rower.find(class_ = 'blockInfo__title').get_text().strip()

                        block_info = tb.find_all(class_='blockInfo__section')
                        for bs in block_info:
                            cont_sec_title = bs.find_all(class_='section__title')
                            for its in cont_sec_title:
                                cont_main_title_mass.append(its.get_text().strip().replace('\n', ''))
                            cont_sec_info = bs.find_all(class_='section__info')
                            for its in cont_sec_info:
                                try:
                                    titls_table_main.append(
                                        {cont_main_title_mass[j]: its.get_text().replace('\n', '').replace('\xa0',
                                                                                                            ' ').strip()})
                                except IndexError:
                                    titls_table_main.append(
                                        {'': its.get_text().replace('\n', '').replace('\xa0', ' ').strip()})

                                j = j + 1
                            conteiner_titls_data[tit] = titls_table_main
                        t = j
                        #Поиск всех элементов таблицы
                        tablos = tb.find_all('table')
                        for tbn in tablos:
                            c_th = tbn.find_all('th')
                            for th_items in c_th:
                                th_text = th_items.get_text().strip()
                                th_mass.append(th_text)

                            c_td = tbn.find_all('td')
                            for td_items in c_td:
                                td_text = td_items.get_text().strip()
                                try:
                                    conteiner_text.append({th_mass[i]: td_text})
                                except Exception:
                                    i = 0
                                    conteiner_text.append({'': td_text})
                                i += 1

                            table_data[cont_main_title_mass[t]] = conteiner_text
                            t = t + 1
                            conteiner_text = []
            collaps_data.update(conteiner_titls_data)
            collaps_data.update(table_data)

        
        return collaps_data
    except Exception:
        status = 'Произошла ошибка с парсингом Выпадающих элементов'
print(collapse_element('0319300010122000854'))
        

def documents(num):
        block_info_title = []
        infos = []
        sec_attrib = []
        sec_value = []
        col_data = []
        data= {}
        font = []
        files_links = {}
        i = 0
        HEADERS = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.0.0'
                          ' YaBrowser/23.3.0.2246 Yowser/2.5 Safari/537.36', 'accept': '*/*'}
        link = f'https://zakupki.gov.ru/epz/order/notice/ea20/view/documents.html?regNumber={num}'
        req = requests.get(url=link, headers=HEADERS)
        src = req.text
        soup = BeautifulSoup(src, "lxml")
        try:
            col_sm_12 = soup.find_all(class_ ='col-sm-12 blockInfo')
            for col in col_sm_12:
                titles = col.find(class_='blockInfo__title').get_text().strip()
                block_info_title.append(titles)

                try:
                    section_value = col.find_all(class_='section__value docName')
                    for sec in section_value:
                        infos.append(sec.get_text().strip())

                    #######################LINKS##########################
                    link_of_files = col.find_all(class_='blockFilesTabDocs')
                    for lux in link_of_files:
                        luxit = lux.find_all('a')
                        try:
                            for links in luxit:
                                if 'download'  in links.get('href').lower():
                                    linkl = links.get('href')
                                    titk = links.get('title')
                                    files_links[titk] = linkl
                            if len(files_links) > 0:
                                try:
                                    # path = os.path.join(self.main_directory, titles)
                                    os.makedirs(f'{titles}', exist_ok=True)
                                    # os.makedirs(path)
                                except Exception:
                                    pass
                                for title, url in files_links.items():
                                    response = requests.get(url, headers=HEADERS)
                                    with open(f'{titles}/{title}', "wb") as f:
                                        f.write(response.content)
                                files_links = {}
                        except Exception:
                            if len(files_links) > 0:
                                # if not os.path.exists(self.main_directory):
                                # path = os.path.join(self.main_directory, titles)
                                # os.makedirs(path)
                                os.makedirs(f'{titles}', exist_ok=True)
                                for title, url in files_links.items():
                                    response = requests.get(url, headers=HEADERS)
                                    with open(f'{titles}/{title}', "wb") as f:
                                        f.write(response.content)
                                files_links = {}

                    #############################################################
                    col_sm = col.find_all(class_='col-sm')
                    for col_12 in col_sm:
                        sec_values = col_12.find_all(class_='section__value')
                        for val in sec_values:
                            sec_value.append(val.get_text().strip())
                        sec_attributs = col_12.find_all(class_='section__attrib')
                        for col_12 in sec_attributs:
                            try:
                                sec_attrib.append(col_12.get_text().strip().replace('\n', ''))
                                i = i + 1
                            except Exception:
                                i = 0
                                sec_attrib.append( col_12.get_text().strip().replace('\n', ''))
                    front = col.find_all('div', attrs={'style': 'font-size: 14px'})
                    if len(front) > 0:
                        for f in front:
                            font.append(f.get_text().strip())

                    if len(col_sm) == 0:
                        perm = col.find(class_='section__title').get_text().strip()
                        # font = []
                        sec_attrib.append(perm)

                    infos.append(sec_attrib)
                    infos.append(font)
                    data[titles] = infos
                    infos = []
                    font = []
                    sec_value = []
                    sec_attrib = []

                except Exception:
                    pass
                    sec_titl = col.find_all(class_ = 'section__title')
                    for ttl in sec_titl:
                        col_data.append(ttl.get_text().strip())
                    data[titles] = col_data
                    infos = []
                    sec_value = []
                    col_data = []

            status = 'успешный парсинг Документов'

        except Exception:
           status = 'Ошибка парсинга Документов'

        # return block_info_title
        return data
# print(documents('0319300010122000854'))
# global_dict = {}
# global_dict.update(documents('0319300010122000854'))
# try:
#     doc = Document()
#     # for item in global_dict:
#     #     doc.add_paragraph(item)
#     # doc.save(f"{self.main_directory}/Все данные о закупке №{self.num}.docx")
#     # Добавляем заголовок
#     doc.add_heading('Данные о закупке', level=1)

#     # Перебираем данные и добавляем их в документ
#     for key, value in global_dict.items():
#         doc.add_heading(key, level=2)
#         if isinstance(value, list):
#             for item in value:
#                 if isinstance(item, dict):
#                     for sub_key, sub_value in item.items():
#                         doc.add_paragraph(f'{sub_key}: {sub_value}')
#                 else:
#                     doc.add_paragraph(str(item))
#         else:
#             doc.add_paragraph(str(value))
    
#     # Сохраняем файл
#     doc.save("Все данные о закупке.docx")
#     status = 'Файл успешно сохранен'

# except Exception as e:
#     print(f'Ошибка записи файлов: {str(e)}')